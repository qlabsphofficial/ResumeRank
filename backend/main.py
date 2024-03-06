from fastapi import FastAPI, Depends, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from database import SessionLocal, engine, Base
from pydantic import BaseModel
from typing import Optional, BinaryIO
from models import User, JobPosting, Resume, JobApplication, Notification, Experience, Certification
# import datetime as dt
from datetime import datetime
from fastapi.responses import FileResponse
import os
from docx import Document
from docx.shared import Inches
import tempfile

IMAGEDIR = "images/"
FILESDIR = "files/"

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=['*'],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

Base.metadata.create_all(bind=engine)

def get_database():
    db = None

    try:
        db = SessionLocal()
    finally:
        yield db
        db.close()

    
class UserModel(BaseModel):
    username: str
    password: str
    firstname : str
    middlename :str
    lastname :str
    email: str
    contact_no : str
    address : str

    class Config:
        orm_mode = True


class ProfileModel(BaseModel):
    username: str = Form(...)
    firstname : str = Form(...)
    lastname : str = Form(...)
    profile_picture: UploadFile = Form(...)


class JobPostingModel(BaseModel):
    title: str
    job_title: str
    description: str
    post_status: bool
    date_expired : str


class ExperienceModel(BaseModel):
    job_title: str
    company: str
    tenure_start: str
    tenure_end: str

    class Config:
        orm_mode = True


class CertificationModel(BaseModel):
    title: str
    training_center : str
    date: str
    attachment : UploadFile = Form(...)

    class Config:
        orm_mode = True
    

class ResumeModel(BaseModel):
    resume_owner: int
    ed_1: str
    ed_2: str
    ed_3: str
    certifications : list[CertificationModel]
    experiences : list[ExperienceModel]
    summary: str
    ref_1: str
    ref_2: str
    ref_3: str

    class Config:
        orm_mode = True


class JobPostingID(BaseModel):
    id: int


class NotificationModel(BaseModel):
    message : str
    sent_to : int


@app.get('/show_users')
async def show_users(db: Session = Depends(get_database)):
    try:
        all_users = db.query(User).all()
        return { 'response': 'User Retrieval Success', 'users': all_users, 'status_code': 200 }
    except:
        return { 'response': 'User Retrieval Failed', 'status_code': 200 }
    

@app.post('/login')
async def login(username: str, password: str, db: Session = Depends(get_database)):
    try:
        existing_user = db.query(User).filter(User.username == username).first()
        
        if existing_user:
            if existing_user.password == password:
                return { 'response': 'Login successful.', 'user_data': existing_user, 'status_code': 200 }
            else:
                return { 'response': 'Login failed.', 'status_code': 403 }
    except:
        return { 'response': 'Login failed.', 'status_code': 403 }


@app.post('/register')
async def register(user: UserModel, db: Session = Depends(get_database)):
    try:
        existing_user = db.query(User).filter(User.username == user.username).first()

        if not existing_user:
            new_user = User()
            new_user.username = user.username
            new_user.password = user.password
            new_user.firstname = user.firstname
            new_user.middlename = user.middlename
            new_user.lastname = user.lastname
            new_user.email = user.email
            new_user.contact_no = user.contact_no
            new_user.address = user.address
            new_user.profile_picture = ''
            db.add(new_user)
            db.commit()

            newly_added = db.query(User).filter(User.username == user.username).first()

            new_resume = Resume()
            new_resume.resume_owner = newly_added.id
            new_resume.ed_1 = ''
            new_resume.ed_2 = ''
            new_resume.ed_3 = ''
            new_resume.summary = ''
            new_resume.ref_1 = ''
            new_resume.ref_2 = ''
            new_resume.ref_3 = ''
            db.add(new_resume)
            db.commit()
            
            return { 'response': 'Registration successful.', 'status_code': 200 }
        else:
            return { 'response': 'User already exists.', 'status_code': 403 }
    except:
        return { 'response': 'Registration failed.', 'status_code': 400 }
    

@app.post('/edit_profile')
async def edit_profile(form_data: ProfileModel = Depends(), db: Session = Depends(get_database)):
    try:
        # existing_user = db.query(User).filter(User.username == profile.username).first()

        file_path = f"{IMAGEDIR}{form_data.profile_picture.filename}"

        with open(file_path, "wb") as f:
            contents = await form_data.profile_picture.read()
            f.write(contents)

        existing_user = db.query(User).filter(User.username == form_data.username).first()
        if existing_user:
            existing_user.firstname = form_data.firstname
            existing_user.lastname = form_data.lastname
            existing_user.profile_picture = f"{IMAGEDIR}{form_data.profile_picture.filename}"
            db.commit()
            
        return { 'response': 'Update profile successful.', 'status_code': 200 }
    except:
        return { 'response': 'update profile failed.', 'status_code': 400 }


@app.get('/retrieve_user_data')
async def retrieve_dashboard_data(user_id: int, db: Session = Depends(get_database)):
    try:
        user = db.query(User).filter(User.id == user_id).first()

        payload = {}
        payload.update({ 'user_data': user })

        return { 'payload': payload, 'status_code': 200 }
        
    except:
        return { 'response': 'Error retrieving data.', 'status_code': 400 }

    
@app.post('/submit_resume')
async def submit_resume(resume: ResumeModel, db: Session = Depends(get_database)):
    # try:
        existing_resume = db.query(Resume).filter(Resume.resume_owner == resume.resume_owner).first()

        if not existing_resume:
            new_resume = Resume()
            new_resume.resume_owner = resume.resume_owner
            new_resume.ed_1 = resume.ed_1
            new_resume.ed_2 = resume.ed_2
            new_resume.ed_3 = resume.ed_3
            new_resume.summary = resume.summary
            new_resume.ref_1 = resume.ref_1
            new_resume.ref_2 = resume.ref_2
            new_resume.ref_3 = resume.ref_3
            db.add(new_resume)
            db.commit()

        else:
            existing_resume.resume_owner = resume.resume_owner
            existing_resume.ed_1 = resume.ed_1
            existing_resume.ed_2 = resume.ed_2
            existing_resume.ed_3 = resume.ed_3
            existing_resume.summary = resume.summary
            existing_resume.ref_1 = resume.ref_1
            existing_resume.ref_2 = resume.ref_2
            existing_resume.ref_3 = resume.ref_3
            db.commit()


        for exp in resume.experiences:
            existing_experience = db.query(Experience).filter(Experience.resume_id == resume.resume_owner, Experience.job_title == exp.job_title, Experience.company == exp.company).first()

            if not existing_experience:
                new_experience = Experience()
                new_experience.job_title=exp.job_title
                new_experience.company=exp.company
                new_experience.tenure_start = datetime.strptime(exp.tenure_start, '%Y-%m-%d')
                new_experience.tenure_end = datetime.strptime(exp.tenure_end, '%Y-%m-%d')
                new_experience.resume_id=resume.resume_owner
                db.add(new_experience)

            db.commit()


        for certs in resume.certifications:
            existing_certification = db.query(Certification).filter(Certification.resume_id == resume.resume_owner, Certification.title == certs.title, Certification.training_center == certs.training_center).first()

            if not existing_certification:
                new_certification = Certification()
                new_certification.title=certs.title
                new_certification.date=certs.date
                new_certification.attachment= f"{FILESDIR}{certs.attachment}"
                new_certification.training_center= certs.training_center
                new_certification.resume_id=resume.resume_owner
                db.add(new_certification)

            db.commit()

        return { 'response': 'resume submitted', 'status_code': 200 }
    # except:
    #     return { 'response': 'Error retrieving data.', 'status_code': 400 }


@app.get('/show_resumes')
async def show_resumes(db: Session = Depends(get_database)):
    try:
        all_resumes = db.query(Resume).all()
        return { 'response': 'resumes retrieved', 'resumes': all_resumes, 'status_code': 200 }
    except:
        return { 'response': 'resumes Retrieval Failed', 'status_code': 200 }
    

@app.get('/retrieve_resume_data')
async def retrieve_resume_data(user_id: int, db: Session = Depends(get_database)):
    try:
        resume = db.query(Resume).filter(Resume.resume_owner == user_id).first()
        experiences = db.query(Experience).filter(Experience.resume_id == resume.id).first()
        certifications = db.query(Certification).filter(Certification.resume_id == resume.id).first()

        return { 
            'response': 'resume retrieved', 
            'resume': resume, 
            'experiences': experiences, 
            'certifications': certifications, 
            'status_code': 200
        }
    except:
        return { 'response': 'resume Retrieval Failed', 'status_code': 200 }


@app.get('/export_resume_to_word')
async def export_resume_to_word(user_id: int, db: Session = Depends(get_database)):
    # try:
        user, resume = db.query(User, Resume).join(Resume, User.id == Resume.resume_owner).filter(User.id == user_id).first()
        experience = db.query(Experience).filter(Experience.resume_id == user_id).all()
        certification = db.query(Certification).filter(Certification.resume_id == user_id).all()

        doc = Document()
        doc.add_heading(f'{user.firstname} {user.middlename} {user.lastname}', 1)
        doc.add_paragraph(f'{user.email} • {user.contact_no} • @{user.firstname}.{user.lastname}')

        doc.add_heading(f'SUMMARY', 2)
        doc.add_paragraph(f'{resume.summary}')

        doc.add_heading(f'WORK EXPERIENCE', 3)

        for exp in experience:
            doc.add_heading(f'{exp.job_title}                                                                                                                 {exp.tenure_start} - {exp.tenure_end}', 3)
            doc.add_paragraph(f'{exp.company}')

        doc.add_heading(f'CERTIFICATION', 3)
        
        for certs in certification:
            doc.add_heading(f'{certs.title}                                                                                                                                      {certs.date}', 3)
            doc.add_paragraph(f'{certs.training_center}')

        # Save the document
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            doc.save(tmp_file.name)
            tmp_file_path = tmp_file.name

        # Return the generated DOCX file as a downloadable attachment
        return FileResponse(tmp_file_path, filename=f"Resume.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        
    # except:
    #     return { 'response': 'resume Retrieval Failed', 'status_code': 200 }


@app.post('/create_job_posting')
async def create_job_posting(job: JobPostingModel, db: Session = Depends(get_database)):
    try:
        date_expired = datetime.strptime(job.date_expired, "%Y-%m-%d")
        
        new_job = JobPosting()
        new_job.job_title = job.job_title
        new_job.description = job.description
        new_job.post_status = job.post_status
        new_job.date_expired = date_expired

        db.add(new_job)
        db.commit()

        return { 'response': 'job created', 'status_code': 200 }
    except:
        return { 'response': 'Error retrieving data.', 'status_code': 400 }
    

@app.post('/delete_job_posting')
async def delete_job_posting(job: JobPostingID, db: Session = Depends(get_database)):
    try:
        retrieved_job_posting = db.query(JobPosting).filter(JobPosting.id == job.id).first()
        
        if retrieved_job_posting:
            db.delete(retrieved_job_posting)
            db.commit()

        return { 'response': 'job posting deleted.'}
    except:
        return {'response': 'failed to job posting.'}
    

@app.get('/show_jobs')
async def show_jobs(db: Session = Depends(get_database)):
    try:
        all_jobs = db.query(JobPosting).all()
        return { 'response': 'jobs retrieved', 'jobs': all_jobs, 'status_code': 200 }
    except:
        return { 'response': 'User Retrieval Failed', 'status_code': 200 }
    

@app.post('/apply_to_job')
async def apply_to_job(user_id: int, job_id: int, db: Session = Depends(get_database)):
    try:
        job_id = db.query(JobPosting).filter(JobPosting.id == job_id).first()
        resume = db.query(Resume).filter(Resume.resume_owner == user_id).first()

        new_application = JobApplication()
        new_application.resume = resume.id
        new_application.job = job_id.id

        db.add(new_application)
        db.commit()
        return { 'response': 'applied to job', 'status_code': 200 }
    except:
        return { 'response': 'Error retrieving data.', 'status_code': 400 }


@app.get('/show_applications')
async def show_applications(db: Session = Depends(get_database)):
    try:
        all_applications = db.query(JobApplication).all()
        return { 'response': 'applications retrieved', 'applications': all_applications, 'status_code': 200 }
    except:
        return { 'response': 'applications Retrieval Failed', 'status_code': 200 }
    

@app.get('/analyze_resumes')
async def analyze_resumes(job_id: int, db: Session = Depends(get_database)):
    # try:
        job = db.query(JobPosting).filter(JobPosting.id == job_id).first()
        all_applications = db.query(JobApplication).filter(JobApplication.job == job_id).all()
        job_desc = job.description.split()

        top_applicants = []
        
        for application in all_applications:
            resume = db.query(Resume).join(User).filter(Resume.id == application.resume).filter(User.id == Resume.resume_owner).first()
            
            resume_analysis = ''
            
            current_points = 0

            if resume.ed_1 and resume.ed_2 and resume.ed_3:
                current_points += 50
            elif resume.ed_1 and resume.ed_2 or resume.ed_1 and resume.ed_3 or resume.ed_2 and resume.ed_3:
                current_points += 25
            elif resume.ed_1 or resume.ed_2 or resume.ed_3:
                current_points += 10

            resume_analysis += f'{resume.summary}'

            resume_text = resume_analysis.split()

            common_words = set(job_desc) & set(resume_text)
            current_points += len(common_words)

            experiences = db.query(Experience).filter(Experience.resume_id == resume.resume_owner).all()

            for experience in experiences:
                experience_analysis = ''

                experience_analysis += f'{experience.job_title}'
                # experience_analysis += f'{experience.company}'

                difference_in_years = (experience.tenure_end - experience.tenure_start).days // 365

                experience_text = experience_analysis.split()
                common_words = set(job_desc) & set(experience_text)

                if len(common_words) > 0:
                    current_points += difference_in_years * 50
                else:
                    current_points += difference_in_years * 20
                
            certifications = db.query(Certification).filter(Certification.resume_id == resume.resume_owner).all()

            for certification in certifications:
                certification_analysis = ''

                certification_analysis += f'{certification.title}'
                # certification_analysis += f'{certification.training_center}'

                certification_text = certification_analysis.split()
                common_words = set(job_desc) & set(certification_text)

                if len(common_words) > 0:
                    current_points += 50
                else:
                    current_points += 20
                
                if certification.attachment:
                    current_points += 10

            if current_points > 250:
                applicant = db.query(User).filter(User.id == resume.resume_owner).first()
                print(applicant.__dict__)
                
                top_applicants.append({'applicant': applicant, 'applicant_resume': resume, 'applicant_points': current_points })

        sorted_top_applicants = sorted(top_applicants, key=lambda x: x['applicant_points'], reverse=True)
        if sorted_top_applicants:
            return { 'response': 'applications retrieved', 'job': job, 'all_applications': all_applications, 'analysis': sorted_top_applicants, 'status_code': 200 }
        # If there are no top applicants, check if there are any applicants at all
        elif all_applications:
            return { 'response': 'no top applicants', 'job': job, 'all_applications': all_applications, 'status_code': 200 }
        # If there are no applicants at all
        else:
            return { 'response': 'no applicants', 'job': job, 'status_code': 200 }
    # except:
    #     return { 'response': 'applications Retrieval Failed', 'status_code': 200 }


@app.post('/create_notification')
async def create_notification(notification: NotificationModel, db: Session = Depends(get_database)):
    try:
        new_notification = Notification()
        new_notification.message = notification.message
        new_notification.sent_to = notification.sent_to

        db.add(new_notification)
        db.commit()

        return { 'response': 'notification created', 'status_code': 200 }
    except:
        return { 'response': 'Error retrieving data.', 'status_code': 400 }
    

@app.get('/show_notifications')
async def get_notifications(id: int, db: Session = Depends(get_database)):
    try:
        all_notifications = db.query(Notification).filter(Notification.sent_to == id).all()

        return {'response': 'retrieval complete.', 'notifications': all_notifications}
    except:
        print('Retrieval failed.')
        return {'response': 'retrieval failed.', 'notifications': all_notifications}
    

@app.post('/delete_notifications')
async def delete_notifications(notification_id: int, db: Session = Depends(get_database)):
    try:
        delete_notification = db.query(Notification).filter(Notification.id == notification_id).first()
        
        if delete_notification:
            db.delete(delete_notification)
            db.commit()

        return { 'response': 'notification deleted.'}
    except:
        return {'response': 'failed to delete notification.'}